import os
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_codigo_18_digitos(producto, cantidad, peso):
    # Formato: Producto (7) + Cantidad (6) + Peso (5)
    prod_str = str(producto).zfill(7)
    cant_str = str(cantidad).zfill(6)
    peso_str = str(peso).zfill(5)
    return f"{prod_str}{cant_str}{peso_str}"

def crear_imagen_codigo(codigo, filename):
    # Usar Code128 para soportar los 18 dígitos
    code128 = barcode.get_barcode_class('code128')
    # Generar el código de barras sin el texto debajo para mayor claridad, lo agregaremos en el Word
    writer = ImageWriter()
    writer.set_options({
        'write_text': False,
        'module_width': 0.3,
        'module_height': 10.0,
        'quiet_zone': 2.0
    })
    
    # Asegurar que el directorio temporal exista
    os.makedirs('temp_barcodes', exist_ok=True)
    filepath = f"temp_barcodes/{filename}"
    
    # barcode automáticamente añade .png si usamos ImageWriter
    code = code128(codigo, writer=writer)
    saved_path = code.save(filepath)
    return saved_path

def main():
    doc = Document()
    
    # Título
    title = doc.add_heading('Códigos de Barras para Pruebas - Revisión de Órdenes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Pedido de prueba: P-123456\n')
    doc.add_paragraph('Estructura del código de barras (18 dígitos):')
    doc.add_paragraph('• Producto: 7 Dígitos (pos 1-7)')
    doc.add_paragraph('• Cantidad: 6 Dígitos (pos 8-13)')
    doc.add_paragraph('• Peso: 5 Dígitos (pos 14-18)\n')
    
    escenarios = [
        {
            'titulo': 'Escenario 1: Revisión Exacta (Sin discrepancia)',
            'desc': 'Producto: FILTRO DE ACEITE | Requerido: 3 | Escanear: 1 vez',
            'producto': 486600,
            'cantidad': 3,
            'peso': 450,
            'tipo': '18_digitos'
        },
        {
            'titulo': 'Escenario 2: Faltante (Conteo < Requerido)',
            'desc': 'Producto: BUJÍA NGK | Requerido: 4 | Escanear: 2 veces (simulando faltan 2)',
            'producto': 1365800,
            'cantidad': 1,
            'peso': 120,
            'tipo': '18_digitos'
        },
        {
            'titulo': 'Escenario 3: Sobrante (Conteo > Requerido)',
            'desc': 'Producto: PASTILLA DE FRENO | Requerido: 2 | Escanear: 3 veces (simulando sobra 1)',
            'producto': 1394001,
            'cantidad': 1,
            'peso': 800,
            'tipo': '18_digitos'
        },
        {
            'titulo': 'Escenario 4: Producto Incorrecto (Código ajeno al pedido)',
            'desc': 'Producto: AMORTIGUADOR (Ajeno) | Requerido: 0 | Escanear: 1 vez para disparar alerta',
            'producto': 9999999,
            'cantidad': 1,
            'peso': 2100,
            'tipo': '18_digitos'
        },
        {
            'titulo': 'Escenario 5: Producto Misceláneo (Código 7 dígitos)',
            'desc': 'Producto: MISCELÁNEO TORNILLOS | Requerido: 10 | Escanear: 1 vez para abrir modal manual',
            'codigo_directo': '3847561',
            'tipo': '7_digitos'
        },
        {
            'titulo': 'Escenario 6: Código Inválido',
            'desc': 'Producto: DESCONOCIDO | Requerido: 0 | Escanear: 1 vez para disparar alerta (Sonido 1)',
            'codigo_directo': '12345',
            'tipo': 'invalido'
        }
    ]
    
    for i, esc in enumerate(escenarios):
        doc.add_heading(esc['titulo'], level=2)
        doc.add_paragraph(esc['desc'])
        
        if esc['tipo'] == '18_digitos':
            codigo = generar_codigo_18_digitos(esc['producto'], esc['cantidad'], esc['peso'])
            texto_explicativo = f"Código: {codigo} (Prod: {str(esc['producto']).zfill(7)} | Cant: {str(esc['cantidad']).zfill(6)} | Peso: {str(esc['peso']).zfill(5)})"
        else:
            codigo = esc['codigo_directo']
            texto_explicativo = f"Código: {codigo}"
            
        img_path = crear_imagen_codigo(codigo, f"barcode_{i}")
        
        # Añadir imagen centrada
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(img_path, width=Inches(3.5))
        
        # Añadir texto del código debajo
        p_text = doc.add_paragraph(texto_explicativo)
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("-" * 50)
        
    doc.save('Productos_Prueba_Revision.docx')
    print("Documento Word generado exitosamente: Productos_Prueba_Revision.docx")

if __name__ == '__main__':
    main()
