#!/usr/bin/env python3
"""
Genera el documento Word de simulación de escaneo para el prototipo
de Revisión de Órdenes.

Estructura de la etiqueta de 18 dígitos:
  Producto  -> 7 dígitos  (pos 1-7)
  Cantidad  -> 6 dígitos  (pos 8-13)
  Peso      -> 5 dígitos  (pos 14-18)
"""

import os
import io
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ─────────────────────────────────────────────────────────────────────────────
# DATOS DEL PEDIDO (deben coincidir exactamente con app.js)
# ─────────────────────────────────────────────────────────────────────────────
PEDIDO_ID = "123456"

PRODUCTOS = [
    {"sku": "0486600", "nombre": "FILTRO DE ACEITE PREMIUM",                 "cant_pedido": 31},
    {"sku": "1365800", "nombre": "BUJÍA NGK ESTÁNDAR",                       "cant_pedido":  7},
    {"sku": "1394001", "nombre": "CINTA AISLANTE AMARILLO 20 PIES",          "cant_pedido": 10},
    {"sku": "2389475", "nombre": "PASTILLA DE FRENO TRASERA",                "cant_pedido":  8},
    {"sku": "9182736", "nombre": "ACEITE MOTOR 5W-30 SINTÉTICO",             "cant_pedido":  9},
]

AJENO = {"sku": "9999999", "nombre": "PRODUCTO AJENO AL PEDIDO (no pertenece al pedido 123456)", "cant_pedido": 0}

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def construir_etiqueta(sku: str, cantidad: int, peso: int = 0) -> str:
    """Construye la etiqueta de 18 dígitos."""
    p = sku.zfill(7)[:7]
    c = str(cantidad).zfill(6)[:6]
    w = str(peso).zfill(5)[:5]
    return p + c + w

def generar_barcode_png(codigo: str, tmp_dir: str, nombre: str) -> str:
    """Genera un PNG de código de barras Code128 y devuelve la ruta."""
    CODE128 = barcode.get_barcode_class('code128')
    writer = ImageWriter()
    bc = CODE128(codigo, writer=writer)
    options = {
        'module_height': 12.0,
        'module_width': 0.28,
        'quiet_zone': 4.0,
        'font_size': 9,
        'text_distance': 4.0,
        'background': 'white',
        'foreground': 'black',
        'write_text': True,
    }
    path_sin_ext = os.path.join(tmp_dir, nombre)
    bc.save(path_sin_ext, options=options)
    return path_sin_ext + '.png'

def set_cell_bg(cell, hex_color: str):
    """Pone color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    """Agrega bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading_row(table, text: str, cols: int, bg: str, font_color: str = "FFFFFF"):
    """Agrega una fila de encabezado de sección que ocupa todas las columnas."""
    row = table.add_row()
    cell = row.cells[0]
    for i in range(1, cols):
        cell.merge(row.cells[i])
    cell.text = text
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(font_color)
    set_cell_bg(cell, bg)
    return row

# ─────────────────────────────────────────────────────────────────────────────
# CONSTRUIR ESCENARIOS
# ─────────────────────────────────────────────────────────────────────────────

TMP = "/home/ubuntu/revision-prototipo/temp_barcodes"
os.makedirs(TMP, exist_ok=True)

# Cada escenario: dict con campos para el doc
escenarios = []

# ── SECCIÓN 1: Escaneos UNITARIOS (cantidad = 1 por escaneo) ─────────────────
for prod in PRODUCTOS:
    etiqueta = construir_etiqueta(prod["sku"], 1, 0)
    escenarios.append({
        "seccion":     "ESCANEO UNITARIO — 1 pieza por escaneo",
        "sec_color":   "1B3892",
        "tipo":        "Unitario",
        "tipo_color":  "E3F2FD",
        "tipo_txt":    "1A237E",
        "sku":         prod["sku"],
        "nombre":      prod["nombre"],
        "cant_pedido": prod["cant_pedido"],
        "cant_etiq":   1,
        "etiqueta":    etiqueta,
        "nota":        f"Escanear {prod['cant_pedido']} veces para completar el pedido",
        "png_name":    f"unit_{prod['sku']}",
    })

# ── SECCIÓN 2: Escaneos TOTALES (cantidad = total del pedido) ─────────────────
for prod in PRODUCTOS:
    etiqueta = construir_etiqueta(prod["sku"], prod["cant_pedido"], 0)
    escenarios.append({
        "seccion":     "ESCANEO TOTAL — cantidad completa del pedido en un solo escaneo",
        "sec_color":   "1B6B3A",
        "tipo":        "Total completo",
        "tipo_color":  "E8F5E9",
        "tipo_txt":    "1B5E20",
        "sku":         prod["sku"],
        "nombre":      prod["nombre"],
        "cant_pedido": prod["cant_pedido"],
        "cant_etiq":   prod["cant_pedido"],
        "etiqueta":    etiqueta,
        "nota":        "Un solo escaneo completa el pedido",
        "png_name":    f"total_{prod['sku']}",
    })

# ── SECCIÓN 3: Escaneo PARCIAL (cantidad = mitad del pedido) ──────────────────
for prod in PRODUCTOS:
    parcial = max(1, prod["cant_pedido"] // 2)
    etiqueta = construir_etiqueta(prod["sku"], parcial, 0)
    escenarios.append({
        "seccion":     "ESCANEO PARCIAL — cantidad menor a la requerida (genera discrepancia)",
        "sec_color":   "B45309",
        "tipo":        "Parcial",
        "tipo_color":  "FFF8E1",
        "tipo_txt":    "7B3F00",
        "sku":         prod["sku"],
        "nombre":      prod["nombre"],
        "cant_pedido": prod["cant_pedido"],
        "cant_etiq":   parcial,
        "etiqueta":    etiqueta,
        "nota":        f"Registra {parcial} de {prod['cant_pedido']} → genera faltante",
        "png_name":    f"parcial_{prod['sku']}",
    })

# ── SECCIÓN 4: Escaneo SOBRANTE (cantidad mayor al pedido) ────────────────────
for prod in PRODUCTOS:
    sobrante = prod["cant_pedido"] + 2
    etiqueta = construir_etiqueta(prod["sku"], sobrante, 0)
    escenarios.append({
        "seccion":     "ESCANEO SOBRANTE — cantidad mayor a la requerida (genera sobrante)",
        "sec_color":   "7B1FA2",
        "tipo":        "Sobrante",
        "tipo_color":  "F3E5F5",
        "tipo_txt":    "4A148C",
        "sku":         prod["sku"],
        "nombre":      prod["nombre"],
        "cant_pedido": prod["cant_pedido"],
        "cant_etiq":   sobrante,
        "etiqueta":    etiqueta,
        "nota":        f"Registra {sobrante} de {prod['cant_pedido']} → genera sobrante",
        "png_name":    f"sobrante_{prod['sku']}",
    })

# ── SECCIÓN 5: Producto AJENO al pedido ──────────────────────────────────────
etiqueta_ajeno = construir_etiqueta(AJENO["sku"], 1, 0)
escenarios.append({
    "seccion":     "PRODUCTO AJENO — código no pertenece al pedido",
    "sec_color":   "B71C1C",
    "tipo":        "Ajeno",
    "tipo_color":  "FFEBEE",
    "tipo_txt":    "7F0000",
    "sku":         AJENO["sku"],
    "nombre":      AJENO["nombre"],
    "cant_pedido": "—",
    "cant_etiq":   1,
    "etiqueta":    etiqueta_ajeno,
    "nota":        "Debe disparar alerta: 'Producto no pertenece a este pedido'",
    "png_name":    f"ajeno_{AJENO['sku']}",
})

# ── SECCIÓN 6: Misceláneo (código de 7 dígitos, sin etiqueta de 18) ───────────
escenarios.append({
    "seccion":     "MISCELÁNEO — código de 7 dígitos (abre modal de cantidad manual)",
    "sec_color":   "00695C",
    "tipo":        "Misceláneo",
    "tipo_color":  "E0F2F1",
    "tipo_txt":    "004D40",
    "sku":         "3847561",
    "nombre":      "MISCELÁNEOS / VARIOS",
    "cant_pedido": 15,
    "cant_etiq":   "N/A",
    "etiqueta":    "3847561",
    "nota":        "Código de 7 dígitos → abre bottom sheet para ingresar cantidad manualmente",
    "png_name":    "misc_3847561",
})

# ─────────────────────────────────────────────────────────────────────────────
# GENERAR IMÁGENES DE CÓDIGOS DE BARRAS
# ─────────────────────────────────────────────────────────────────────────────
print("Generando códigos de barras...")
for esc in escenarios:
    path = generar_barcode_png(esc["etiqueta"], TMP, esc["png_name"])
    esc["barcode_path"] = path
    print(f"  ✓ {esc['png_name']}: {esc['etiqueta']}")

# ─────────────────────────────────────────────────────────────────────────────
# CONSTRUIR EL DOCUMENTO WORD
# ─────────────────────────────────────────────────────────────────────────────
print("\nGenerando documento Word...")
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Portada ───────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("GUÍA DE SIMULACIÓN DE ESCANEO")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1B, 0x38, 0x92)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run(f"Prototipo — Revisión de Órdenes  |  Pedido {PEDIDO_ID}")
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ── Instrucciones ─────────────────────────────────────────────────────────────
instr = doc.add_paragraph()
instr.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = instr.add_run("Cómo usar este documento:")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x1B, 0x38, 0x92)

bullets = [
    "ESCANEAR: Apunta el lector de código de barras al código impreso.",
    "COPIAR/PEGAR: Copia el texto del campo 'Código (texto)' y pégalo en el campo de escaneo del prototipo.",
    "La etiqueta de 18 dígitos tiene estructura: [Producto 7 dígitos][Cantidad 6 dígitos][Peso 5 dígitos].",
    "Los escaneos UNITARIOS suman 1 pieza al conteo; los TOTALES completan el pedido en un solo escaneo.",
    "El producto AJENO debe disparar una alerta de error en el prototipo.",
    "El MISCELÁNEO (7 dígitos) abre el modal de ingreso manual de cantidad.",
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b).font.size = Pt(10)

doc.add_paragraph()

# ── Tabla de resumen del pedido ───────────────────────────────────────────────
resumen_title = doc.add_paragraph()
r = resumen_title.add_run("Artículos del Pedido 123456")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x1B, 0x38, 0x92)

tbl_res = doc.add_table(rows=1, cols=3)
tbl_res.style = 'Table Grid'
hdr = tbl_res.rows[0].cells
hdr[0].text = "SKU"
hdr[1].text = "Nombre"
hdr[2].text = "Cant. Pedido"
for cell in hdr:
    set_cell_bg(cell, "1B3892")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)

for prod in PRODUCTOS:
    row = tbl_res.add_row().cells
    row[0].text = prod["sku"]
    row[1].text = prod["nombre"]
    row[2].text = str(prod["cant_pedido"])
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, "F5F7FF")

# Misceláneo
row = tbl_res.add_row().cells
row[0].text = "3847561"
row[1].text = "MISCELÁNEOS / VARIOS"
row[2].text = "15"
for cell in row:
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cell, "E0F2F1")

doc.add_paragraph()
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TARJETAS DE ESCENARIO
# ─────────────────────────────────────────────────────────────────────────────
secciones_vistas = set()

for esc in escenarios:
    # Encabezado de sección (solo la primera vez)
    if esc["seccion"] not in secciones_vistas:
        secciones_vistas.add(esc["seccion"])
        sec_para = doc.add_paragraph()
        sec_run = sec_para.add_run(f"▌ {esc['seccion']}")
        sec_run.bold = True
        sec_run.font.size = Pt(12)
        # Color según sección
        color_hex = esc["sec_color"]
        r_val = int(color_hex[0:2], 16)
        g_val = int(color_hex[2:4], 16)
        b_val = int(color_hex[4:6], 16)
        sec_run.font.color.rgb = RGBColor(r_val, g_val, b_val)
        doc.add_paragraph()

    # Tabla de tarjeta (2 columnas: info | código de barras)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(10)
    tbl.columns[1].width = Cm(7)

    # Fila 1: tipo + SKU
    row1 = tbl.add_row()
    row1.cells[0].merge(row1.cells[1])
    p = row1.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_tipo = p.add_run(f"  {esc['tipo'].upper()}  ")
    r_tipo.bold = True
    r_tipo.font.size = Pt(9)
    c = esc["tipo_txt"]
    r_tipo.font.color.rgb = RGBColor(int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))
    set_cell_bg(row1.cells[0], esc["tipo_color"])

    # Fila 2: datos + código de barras
    row2 = tbl.add_row()
    cell_info = row2.cells[0]
    cell_bc   = row2.cells[1]
    cell_info.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_bc.vertical_alignment   = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell_info, "FFFFFF")
    set_cell_bg(cell_bc,   "FFFFFF")

    # Info
    def add_info_line(cell, label, value, bold_val=False):
        p = cell.add_paragraph()
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.font.size = Pt(9)
        r_lbl.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        r_val = p.add_run(str(value))
        r_val.font.size = Pt(9)
        r_val.bold = bold_val
        r_val.font.color.rgb = RGBColor(0x1B, 0x38, 0x92) if bold_val else RGBColor(0x21, 0x21, 0x21)
        return p

    # Limpiar párrafo vacío inicial
    cell_info.paragraphs[0].clear()

    add_info_line(cell_info, "SKU",          esc["sku"],         bold_val=True)
    add_info_line(cell_info, "Producto",     esc["nombre"])
    add_info_line(cell_info, "Cant. pedido", esc["cant_pedido"])
    add_info_line(cell_info, "Cant. etiq.",  esc["cant_etiq"])

    # Campo de texto para copiar/pegar
    p_copy = cell_info.add_paragraph()
    r_lbl2 = p_copy.add_run("Código (texto): ")
    r_lbl2.font.size = Pt(9)
    r_lbl2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    r_cod = p_copy.add_run(esc["etiqueta"])
    r_cod.bold = True
    r_cod.font.size = Pt(10)
    r_cod.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # Fondo amarillo para destacar el campo copiable
    from docx.oxml import OxmlElement as OE
    from docx.oxml.ns import qn as QN
    rPr = r_cod._r.get_or_add_rPr()
    highlight = OE('w:highlight')
    highlight.set(QN('w:val'), 'yellow')
    rPr.append(highlight)

    add_info_line(cell_info, "Nota", esc["nota"])

    # Código de barras
    bc_p = cell_bc.paragraphs[0]
    bc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        bc_p.add_run().add_picture(esc["barcode_path"], width=Cm(6.0))
    except Exception as e:
        bc_p.add_run(f"[Error: {e}]")

    doc.add_paragraph()  # espacio entre tarjetas

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run("Documento de uso interno — Prototipo Revisión de Órdenes")
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = "/home/ubuntu/revision-prototipo/Guia_Simulacion_Escaneo.docx"
doc.save(output_path)
print(f"\n✓ Documento guardado: {output_path}")
print(f"  Tamaño: {os.path.getsize(output_path) // 1024} KB")
print(f"  Total de escenarios: {len(escenarios)}")
